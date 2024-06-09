import streamlit as st
import google.generativeai as genai
from PIL import Image
import markdown
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from io import BytesIO
import os

os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
model = genai.GenerativeModel('gemini-1.5-flash-latest')

def response(image):
    prompt = """You are an intelligent document creator. Could you please extract the words from the given screenshot and provide me document text that matches exact screenshot font and look
    important note: if the screenshot not contain any text means you must say 'please upload a valid screenshot'"""
    img = Image.open(image)
    response = model.generate_content([prompt, img]) 
    return response.text

def markdown_to_word(markdown_text):
    # Create a new Word document
    doc = Document()

    for line in markdown_text.split('\n'):
        if line.startswith('# '):
            heading = line[2:]
            p = doc.add_heading(heading, level=1)
        elif line.startswith('## '):
            heading = line[3:]
            p = doc.add_heading(heading, level=2)
        elif line.startswith('### '):
            heading = line[4:]
            p = doc.add_heading(heading, level=3)
        elif line.startswith('- '):
            item = line[2:]
            p = doc.add_paragraph(item, style='ListBullet')
        else:
            p = doc.add_paragraph()
            words = line.split(' ')
            for word in words:
                word = word.strip()
                if word.startswith('**') and word.endswith('**'):
                    run = p.add_run(word[2:-2])
                    run.bold = True
                elif word.startswith('*') and word.endswith('*'):
                    run = p.add_run(word[1:-1])
                    run.italic = True
                else:
                    p.add_run(word)
                p.add_run(' ')

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.title("SCREENSHOT🖼️ - DOCUMENT📃")
st.markdown("""
    <style>
    .justified-text {
        text-align: justify;
    }
    </style>
    """, unsafe_allow_html=True)
with st.sidebar:
    st.header("ABOUT:")
    
    st.caption("""
        <div class="justified-text">
            Screenshot to Document file Creator is an AI powered app that allows users to effortlessly convert their screenshots into Word documents. Simply upload a screenshot, and the app will generate a Word document based on the image provided, ensuring a seamless and efficient conversion process. Ideal for anyone looking to quickly turn visual content into editable text documents.
        </div>
        """, unsafe_allow_html=True)
    
    for _ in range(17):
        st.write("") 
    st.subheader("Build By:")
    st.write("[Pachaiappan❤️](https://mr-vicky-01.github.io/Portfolio)")
    st.write("contact: [Email](mailto:pachaiappan1102@gamil.com)")
    
fake_image_text = 'please upload a valid screenshot.'
st.text("Upload your screenshot to convert it into a Word document")
uploaded_file = st.file_uploader("", type=["png", "jpg", "jpeg"])
if uploaded_file:
    st.image(uploaded_file)
    button = st.button("Generate Document")
    if button:
        with st.spinner("Generating a Document..."):
            text = response(uploaded_file)
        st.write(text)
        
        if text.lower().strip() != fake_image_text:
            doc_buffer = markdown_to_word(text)
            st.download_button(
            label="Download",
            data=doc_buffer,
            file_name="output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            